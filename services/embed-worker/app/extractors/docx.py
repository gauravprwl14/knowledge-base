"""DOCX text extractor using python-docx."""

import asyncio
from pathlib import Path

from app.extractors.base import BaseExtractor


class DocxExtractor(BaseExtractor):
    """Extract plain text from Microsoft Word (.docx) files.

    Reads all paragraph text and table cell text from the document,
    joining them with newlines into a single string.
    """

    supported_mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    async def extract(self, file_path: Path) -> str:
        """Extract text content from a .docx file.

        Args:
            file_path: Path to the .docx file on disk.

        Returns:
            All text content from paragraphs and table cells, joined by newlines.

        Raises:
            ExtractionError: If the document cannot be read or parsed.
        """
        return await asyncio.to_thread(self._extract_sync, file_path)

    def _extract_sync(self, file_path: Path) -> str:
        """Synchronous extraction logic, intended for thread pool execution.

        Args:
            file_path: Path to the .docx file.

        Returns:
            Extracted text content as a single string.
        """
        import docx  # python-docx

        doc = docx.Document(str(file_path))
        lines: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)

        return "\n".join(lines)
